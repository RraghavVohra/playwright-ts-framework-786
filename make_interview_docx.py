# -*- coding: utf-8 -*-
"""Generates Playwright-Interview-Prep.docx from in-script content.
Companion to Concepts-Study-Guide.docx (JS/TS language fundamentals).
This one covers Playwright/framework-level interview questions, grounded in
our own playwright-ts-framework repo (config, fixtures, page objects, Fixes.md)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- base styles (identical to Concepts-Study-Guide.docx for visual consistency) ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # dark blue
CODEBG = 'F2F2F2'

def shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p

def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def numbered(text):
    return doc.add_paragraph(text, style='List Number')

def code(text):
    p = doc.add_paragraph()
    shade(p, CODEBG)
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), 'Consolas')
    rfonts.set(qn('w:hAnsi'), 'Consolas')
    return p

def soundbite(text):
    p = doc.add_paragraph()
    shade(p, 'FFF2CC')
    p.paragraph_format.left_indent = Inches(0.1)
    r = p.add_run('Interview soundbite:  ')
    r.bold = True
    r2 = p.add_run(text)
    r2.italic = True
    return p

def fixme(text):
    p = doc.add_paragraph()
    shade(p, 'FCE4D6')
    p.paragraph_format.left_indent = Inches(0.1)
    r = p.add_run('Honest gap / known issue:  ')
    r.bold = True
    r2 = p.add_run(text)
    r2.italic = True
    return p

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, htext in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(htext)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()
    return t

# ============================ TITLE ============================
doc.add_heading('Playwright Interview Prep', level=0)
sub = doc.add_paragraph()
sr = sub.add_run('Real answers grounded in our own playwright-ts-framework repo — '
                  'not textbook theory. Companion to Concepts-Study-Guide.docx.')
sr.italic = True
sr.font.size = Pt(11)
doc.add_paragraph()
para('How to read this: every answer below is tied to something real in this codebase — '
     'config.ts, fixtures.ts, a page object, or an entry in Fixes.md. Where I genuinely '
     "haven't done something yet (e.g. MCP), it's marked honestly instead of faked.", italic=True)
doc.add_page_break()

# ============================ 1. FRAMEWORK TYPE ============================
h1('1. What kind of framework is this? (Data-driven / BDD / keyword-driven)')
para("It's a Page Object Model framework with data-driven elements layered in. It is "
     "NOT BDD (no Cucumber/Gherkin, no .feature files) and NOT keyword-driven "
     "(no generic action-sheet engine reading from Excel/CSV).")
bullet('POM — pages/ holds one class per screen (PushNotificationPage, DocumentLibraryPage, '
       'SocialAutoPostPage, TestimonialsPage).')
bullet('Data-driven where it matters — utils/config.ts drives behaviour per environment '
       '(dev/preprod/prod/digipulse) from .env; social-autopost tests for 13 image sizes are '
       'generated from a data array in a loop instead of copy-pasted.')
soundbite("It's POM-based with data-driven elements — environment config drives which server/"
          "credentials a run uses, and repetitive variations are generated from a data array "
          "rather than duplicated. I didn't go BDD because there's no non-technical stakeholder "
          "consuming Gherkin specs on this project — it would add translation overhead nobody reads.")

# ============================ 2. ASYNC/AWAIT (CONDENSED) ============================
h1('2. async / await / Promise — quick recap')
para('Full deep-dive with diagrams lives in Concepts-Study-Guide.docx, section 3 & 3B. '
     'Condensed version for interview recall:', italic=True)
bullet('async marks a function so it always returns a Promise (even Promise<void> if it returns nothing).')
bullet('await pauses only inside that async function, until the Promise resolves, then unwraps the value.')
bullet('Promise = a JS built-in object representing "a value that will exist later" — not a TS-only concept; '
       'TS just lets you type what that future value will be.')
bullet('Forgetting await does NOT throw an error — the next line just races ahead before the browser '
       'finished. That race condition is the #1 cause of flaky Playwright tests.')
soundbite('Async/await is Promise syntax sugar — await does what .then() does, just written to read '
          'top-to-bottom. Forgetting await in Playwright is silent — no error, just a race condition.')

# ============================ 3. SELENIUM VS PLAYWRIGHT ============================
h1('3. Selenium vs Playwright')
para('Lead with architecture, not a feature checklist — that is the senior-level framing.')
table(['Aspect', 'Selenium', 'Playwright'],
      [['Communication', 'WebDriver protocol — HTTP to a driver process (chromedriver), which talks to the browser', 'Direct CDP/WebSocket connection to the browser — no driver process'],
       ['Waits', 'No built-in actionability checks — manual WebDriverWait + ExpectedConditions', 'Auto-waits before every action: visible, stable, receives events, enabled'],
       ['Speed', 'Slower — extra HTTP hop per command', 'Faster — direct protocol'],
       ['Multi-tab/context', 'getWindowHandles() + switchTo() — painful', 'browser.newContext() — native, isolated sessions'],
       ['API testing', 'Needs RestAssured / separate library', "Built-in 'request' fixture — no browser needed"],
       ['Setup', 'Driver version must match browser version exactly', 'npx playwright install — browsers bundled, versions locked to the Playwright version'],
       ['Debugging', 'Screenshots + manual logging', 'Trace Viewer, Inspector, Codegen, UI mode'],
       ['Network control', 'Needs BrowserMob Proxy or similar', "Built-in page.route() for interception/mocking"]])
soundbite('The core difference is architectural — Selenium goes through a driver process over HTTP, '
          'Playwright talks to the browser directly. That single difference is what makes auto-wait, '
          'multi-context, and network interception native in Playwright instead of bolt-on workarounds.')

# ============================ 4. AI TOOLS ============================
h1('4. AI tools used with Playwright')
para('Answer with what is actually true for this project: solo contributor, using Claude Code as both '
     'pair-programmer and the code reviewer there is no second engineer to provide.')
bullet('Root-cause diagnosis on failures — feed the error + relevant page-object code to Claude and ask '
       '"why would this fail on preprod but not dev." This is literally how Fix 10 (XPath text() vs '
       'normalize-space()) and Fix 8 (dropdown race condition) got solved.')
bullet('Pattern reuse across the codebase — once a fix is found in one page object, ask "where else does '
       'this pattern exist?" and apply consistently (Fix 8 -> Fix 12, same retry pattern in a second page object).')
bullet('Code review on own PRs — no second engineer, so Claude is asked specifically about strict-mode '
       'risk and race conditions before a push.')
bullet('Documentation that compounds — PROGRESS.md and Fixes.md are AI-assisted: every fix becomes a '
       'structured Problem -> Fix -> Interview-talking-point note immediately, so docs double as interview prep.')
soundbite('I work solo on this framework, so I use Claude Code as both a pair programmer and the code '
          'reviewer I do not have — feeding it failing test output and the relevant page object to diagnose '
          'root cause, and asking whether a fix pattern from one page object should be applied elsewhere.')

# ============================ 5. MCP ============================
h1('5. How are you using MCP?')
fixme('@playwright/mcp is not yet wired into this repo (no .mcp.json, no @playwright/mcp in package.json '
      'at time of writing). Plan: set it up for real, use it once on something concrete, then convert this '
      'section into an experience-based answer.')
para('Conceptual answer until then: @playwright/mcp is Microsoft\'s MCP server that exposes browser '
     'actions (navigate, click, snapshot the accessibility tree) as tools an AI agent can call directly — '
     'so the agent drives a live browser itself, useful for exploratory testing or scaffolding a test\'s '
     'first draft from a completed task, instead of writing Playwright code by hand first.')
soundbite("I haven't wired this into the project yet, but I understand it as exposing Playwright's browser "
          "actions as MCP tools so an agent can drive the browser directly — useful for exploratory testing "
          "or a first-draft test scaffold. [Update this once hands-on.]")

# ============================ 6. LOCATORS (SIMPLE) ============================
h1('6. What locators have you used in Playwright?')
para('Keep this simple — four things, one real example each from our own page objects.')
table(['Locator type', 'Real example from this repo', 'When used'],
      [['ID', "#username, #password (auth.setup.ts)", 'Best — use whenever the app has one'],
       ['CSS class', '.ms-selectall.global', 'Library-generated components with no usable ID'],
       ['XPath (text/contains/ancestor)', "//a[normalize-space()='Push Notification']", 'Legacy app with no test-ids; also the ONLY way to walk UP the DOM (ancestor::tr) — CSS cannot do that at all'],
       ["Playwright's own locators (getByRole / getByText)", "getByRole('button', { name: 'Schedule Post' })", 'Newer code — more resistant to markup/class changes than CSS/XPath']])
para('Priority order, said out loud if asked: getByRole -> getByLabel -> getByText -> CSS/ID -> XPath last.')
soundbite('I default to role/text-based locators, but this is a legacy Angular/jQuery UI app without '
          'consistent test-ids, so XPath is still load-bearing — especially ancestor traversal, which '
          'CSS structurally cannot do.')

# ============================ 7. FIXTURES (SIMPLE DEF FIRST) ============================
h1('7. What are fixtures in Playwright?')
h2('Simple definition first')
para('A fixture is a reusable setup-and-teardown function that Playwright runs automatically for a test. '
     'You ask for it by name as a parameter, and Playwright builds it and hands it to you — no "new" needed.')
h2('The mechanism (Dependency Injection)')
para('Under the hood, this is Playwright\'s dependency injection system. You declare what you need as a '
     'parameter name; Playwright matches it to a registered fixture, runs that function, and injects the '
     'result. The fixture function controls setup (before use()) and teardown (after use()).')
code("documentLibraryPage: async ({ page }, use) => {\n"
     "  const documentLibraryPage = new DocumentLibraryPage(page);   // setup\n"
     "  await use(documentLibraryPage);                              // hand to test, test runs here\n"
     "  // anything after use() = teardown, runs even if the test fails\n"
     "}")
para('Our own utils/fixtures.ts is the textbook example — every page object (PushNotificationPage, '
     'DocumentLibraryPage, SocialAutoPostPage, TestimonialsPage) is wired this way, so a test just '
     'declares { documentLibraryPage } and gets a ready instance with zero "new" in the test file.')
soundbite('Fixtures are dependency injection for tests — you declare a need by parameter name, Playwright '
          'resolves it. I use this to inject page objects so test files never manually instantiate classes.')

# ============================ 8. BUILT-IN FIXTURES ============================
h1('8. Playwright\'s built-in fixtures')
table(['Fixture', 'Scope', 'What it gives you'],
      [['page', 'test', 'A fresh browser tab — the one used most often'],
       ['context', 'test', 'The browser context (isolated cookies/storage) that owns page'],
       ['browser', 'worker', 'The raw browser instance — rarely needed directly'],
       ['browserName', 'worker', "String: 'chromium' / 'firefox' / 'webkit' — for conditional logic per browser"],
       ['request', 'worker', 'API testing client — make HTTP calls with no browser at all']])
para('Tie back to our repo: auth.setup.ts uses the built-in page fixture directly (it is the one-time '
     'login step, outside our custom fixtures file), and page.context().storageState(...) reaches the '
     'built-in context fixture from inside a page-based test.')

# ============================ 9. CUSTOM FIXTURE EXAMPLE ============================
h1('9. Example of a customized fixture')
para('Use the real one from utils/fixtures.ts:')
code("type MyFixtures = {\n"
     "  pushNotificationPage: PushNotificationPage;\n"
     "  documentLibraryPage:  DocumentLibraryPage;\n"
     "  socialAutoPostPage:   SocialAutoPostPage;\n"
     "  testimonialsPage:     TestimonialsPage;\n"
     "};\n\n"
     "export const test = base.extend<MyFixtures>({\n"
     "  documentLibraryPage: async ({ page }, use) => {\n"
     "    const documentLibraryPage = new DocumentLibraryPage(page);\n"
     "    await use(documentLibraryPage);\n"
     "  },\n"
     "  // ...same pattern for the other 3\n"
     "});\n\n"
     "export { expect };   // re-exported so test files only import from ONE place")
para('Next-level example if asked "what would you add next": a fixture that depends on ANOTHER custom '
     'fixture, or one with real teardown — e.g. a fixture that creates a test document via API in setup '
     'and deletes it via API in teardown, so tests never leave junk data on preprod. Not built yet — good '
     'honest answer for growth, and it connects directly to the storageState + API-auth combo in section 20.')

# ============================ 10. CONFIG WALKTHROUGH ============================
h1('10. Settings in our playwright.config.ts')
para('Walk through the real file with the WHY for each — strongest possible answer because every line maps to a real decision.')
table(['Setting', 'Our value', 'Why'],
      [['globalSetup', './utils/global-setup', 'Wipes allure-results/ before anything runs — clean slate, cross-platform via fs (no shell-specific rm)'],
       ['testDir', './tests/e2e', 'Scoped so old learning files elsewhere in the repo never get picked up'],
       ['timeout', '90 * 1000', 'Raised from 60s after tests timed out in beforeEach when preprod was slow late in a sequential run (Fix 18)'],
       ['expect.timeout', '15 * 1000', 'Separate, shorter budget for individual assertions vs the whole test'],
       ['fullyParallel', 'false', 'All tests share ONE login session via auth.json — parallel workers would conflict'],
       ['workers', '1', 'Same reason — one browser at a time, deterministic, easier to debug'],
       ['reporter', "['line'], ['allure-playwright'], ['html']", 'line = live terminal feedback, Allure = rich history/trends, HTML = local trace inspection'],
       ['use.trace', "'on-first-retry' (toggled 'on' during active dev)", 'Trade-off between storage cost and debuggability'],
       ['use.screenshot', "'only-on-failure'", "Don't pay the cost on passing tests"],
       ['use.headless', 'false', 'Visible browser locally; CI uses Azure cloud browsers instead'],
       ['projects', 'setup -> chromium (dependencies)', "Two-project login pattern: 'setup' runs auth.setup.ts once, writes auth.json; 'chromium' depends on it and loads storageState"]])
soundbite("Every setting in my config has a reason tied to a real problem I hit — workers: 1 isn't a "
          "default I left alone, it's because parallel workers would race on the same auth.json session. "
          "timeout: 90s is the documented fix for a real flaky-test incident in Fixes.md, not an arbitrary number.")

# ============================ 11. CHALLENGES FACED ============================
h1('11. Challenges faced — three ready answers (pick based on what they probe for)')

h2('A. Flaky test from a race condition (best general-purpose answer)')
para("clickUploadOption() was timing out — the Actions dropdown was closing before Playwright's waitFor "
     "could catch the option visible. I was watching a permanently-hidden element for the full timeout "
     "with no way to recover. Fixed it with expect().toPass() — Playwright's retry wrapper for arbitrary "
     "async blocks — which re-opens the dropdown and re-checks visibility on every retry. I then found "
     "the exact same pattern in a second page object and applied the identical fix there too.")
code("async clickUploadOption(): Promise<void> {\n"
     "  await expect(async () => {\n"
     "    if (!(await this.uploadMenuOption.isVisible())) {\n"
     "      await this.actionsButton.click();\n"
     "    }\n"
     "    await this.uploadMenuOption.waitFor({ state: 'visible', timeout: 2000 });\n"
     "  }).toPass({ timeout: 30000 });\n"
     "  await this.uploadMenuOption.click();\n"
     "}")

h2('B. CI/CD failure that was NOT in the test code at all (best for "passes locally, fails in CI")')
para("Tests passed locally but failed in GitHub Actions with 'value: expected string, got undefined' on "
     "the username field. Root cause: .env is gitignored — correctly, it holds credentials — so "
     "USER_EMAIL/USER_PASSWORD simply did not exist as env vars on the CI runner. Fixed by adding them as "
     "GitHub Secrets, and since the framework supports four environments with different credentials each, "
     "used GitHub Environments so the same secret NAMES resolve to different VALUES depending on which "
     "environment is picked from a manual dropdown — zero code changes needed to onboard a new environment.")

h2('C. CI/CD pipeline efficiency (good if they probe pipeline maturity)')
para("A browser-install step was taking 25+ minutes on every CI run and never benefited from caching. "
     "Root cause: I was manually cancelling runs out of impatience, and actions/cache only persists its "
     "cache in a post-job step that runs on completion — a cancelled run starts every subsequent run from "
     "zero again. Once traced, I also realised the install was unnecessary altogether, since tests run on "
     "Azure's cloud browsers, not a local Chromium — removed the step entirely instead of fixing the cache.")

h2('Solo-dev follow-up (code review / unit testing / CI vs local)')
para("Since I work solo, there is no second engineer for review, so I treat Claude as that reviewer — feed "
     "it the diff, ask specifically about strict-mode risk, race conditions, and whether a fix should be "
     "applied elsewhere. Playwright E2E tests don't have unit tests in the traditional sense, but I run "
     "npx tsc --noEmit as a correctness gate before any push, and dry-run new locators against multiple "
     "environments before trusting them — this app renders the same button differently per environment, "
     "and that habit has caught more real bugs than a unit test would have here.")

# ============================ 12. IFRAMES ============================
h1('12. How do you handle iframes in Playwright?')
fixme('This app does not actually use iframes anywhere in the repo (no frameLocator calls in any page '
      "object). Be honest about that if asked directly, then show the mechanism.")
code("const frame = page.frameLocator('iframe#myframe');\n"
     "await frame.locator('button').click();")
para('Why special handling is needed: page.locator() only sees the main document. An iframe is a '
     'SEPARATE HTML document embedded inside — Playwright is blind to it until you switch context via '
     'frameLocator(). Once you have that frame context, every interaction goes through it.')
soundbite("This app doesn't use iframes, so I haven't needed frameLocator() here — but I'd reach for it "
          "the moment a third-party widget (payment gateway, embedded chat) showed up.")

# ============================ 13. WAITS ============================
h1('13. Different kinds of waits in Playwright')
numbered('Auto-wait — built into every action (click, fill): checks visible/stable/enabled/receives-events.')
numbered('waitFor({ state }) — explicit wait for visible/hidden/attached/detached.')
numbered('waitForNavigation / waitForURL — for page navigations completing.')
numbered('waitForLoadState() — domcontentloaded vs networkidle. We hit the real difference: networkidle '
         'never resolves on preprod because of background polling, so domcontentloaded is used when only '
         'DOM rows are needed, networkidle only when an element genuinely needs full JS settle.')
numbered('waitForSelector — used for delayed validation messages that appear seconds AFTER a click; more '
         'reliable than a bare expect().toBeVisible() for that case.')
numbered('waitForFunction — polls an arbitrary JS condition (used to confirm a datetime input populated '
         'after a calendar day click).')
numbered('waitForTimeout — hard sleep, last resort. We still have one for an xdsoft calendar animation '
         'with no deterministic signal — documented WHY rather than hidden as a smell.')
numbered('expect().toPass() — not technically a wait, but the most powerful one: retries an entire async '
         'block until it passes, and can take CORRECTIVE action (re-click) between retries, not just wait passively.')
soundbite('The one most people forget is expect().toPass() — it is not a wait, it is a retry-with-recovery '
          'loop. A plain wait just watches; toPass() can re-trigger the precondition on every attempt. '
          'That distinction fixed a real flaky test for me (Fix 8).')

# ============================ 14. ALERTS ============================
h1('14. How do you handle alerts in Playwright?')
para('Answer with the Type 1 vs Type 2 distinction — most candidates blur this.')
h2('Type 1 — in-app HTML popups')
para('A styled <div> made to look like a dialog — just a normal locator, no special API. Our delete-'
     "confirmation flow (dialogBox, okButton in DocumentLibraryPage) is exactly this: it's an HTML "
     'element, not a native browser dialog, so page.locator(...).click() handles it like anything else.')
h2('Type 2 — native browser dialogs (window.alert / confirm / prompt)')
para('These live OUTSIDE the DOM — no CSS selector can ever reach them. Need a listener registered '
     'BEFORE the action that triggers them:')
code("page.on('dialog', async dialog => {\n"
     "  console.log(dialog.message());\n"
     "  await dialog.accept();      // or dialog.dismiss(), or dialog.accept('text') for a prompt\n"
     "});\n"
     "await page.locator('#triggerBtn').click();   // action comes AFTER the listener")
soundbite('Most apps mix both — a styled confirmation modal that is just DOM, and an actual window.confirm. '
          'The tell: if DevTools shows it as an element, it is Type 1 and you just click it. If DevTools '
          'shows nothing because it is browser-native, you need page.on(\'dialog\'), registered before the trigger.')

# ============================ 15. FRAMEWORK WALKTHROUGH ============================
h1('15. "Walk me through your framework" — the script')
para('A tight 60-90 second answer. Practice saying this without sounding rehearsed.', italic=True)
code(
"It's a Playwright + TypeScript framework testing a multi-environment web app -- dev, preprod, prod, "
"and a fourth client environment called Digipulse -- switched purely through one .env variable, no code "
"changes.\n\n"
"It follows Page Object Model: four page objects right now -- Push Notification, Document Library, "
"Social Auto-post, Testimonials -- each holding its locators and action methods. Tests never touch raw "
"locators directly.\n\n"
"Page objects are injected into tests through custom fixtures -- Playwright's dependency injection -- so "
"a test just declares { documentLibraryPage } and gets a ready instance.\n\n"
"For auth, I use a two-project setup: a 'setup' project logs in once and saves the session via "
"storageState to auth.json; the main 'chromium' project depends on that and loads the saved session -- "
"every test starts already logged in, login happens once per run, not once per test.\n\n"
"Tests are tagged @smoke and @regression -- one critical happy-path per feature is @smoke, everything is "
"@regression -- so I can run a fast sanity check or the full suite via --grep.\n\n"
"Reporting is three-layered: 'line' for live terminal feedback, Allure for a rich historical report, "
"HTML for local trace debugging.\n\n"
"CI is GitHub Actions, manually triggered via workflow_dispatch with an environment dropdown, running on "
"Azure's cloud browsers rather than installing Chromium on the runner -- cut a 25-minute install step to zero.\n\n"
"Right now it's 50 test cases across the four features. I'm upfront that two Testimonials tests are still "
"failing after I added that feature most recently -- I track that in my fix log rather than hiding it, "
"and it's next on my list."
)
para('That last line is deliberate — naming a known gap unprompted reads as integrity, not weakness, IF '
     'said calmly and followed by "next on my list." A rehearsed-perfect answer is a red flag to good interviewers.', italic=True)

# ============================ 16 & 17. PROCESS ============================
h1('16. "How many test cases do you automate every day?"')
para('No safe fixed number — any number either sounds inflated or low without context. Answer with FACTORS, not a number.')
soundbite("It depends on complexity and locator stability. A simple navigation/validation test might take "
          "15-20 minutes. A full E2E flow with file upload, a third-party date picker, and environment-"
          "conditional steps can take half a day, because I'm also debugging real app behaviour differences "
          "across environments, not just writing Playwright code. Across this framework, roughly 50 tests "
          "over about 13 sessions -- several of those sessions were spent fixing flaky tests and CI issues "
          "rather than writing new ones, and I'd rather report that honestly. A more useful metric than "
          "'tests per day' is how fast I can turn a real production bug into a regression test that "
          "prevents recurrence -- that's what I actually optimise for.")

h1('17. What process do you follow for automation?')
numbered('Understand the requirement / manual test case — actual user flow, edge cases, negative paths.')
numbered('Check for an existing locator/page object — extend, do not duplicate.')
numbered('Identify environment differences early — does this element render differently dev vs prod?')
numbered('Write the page object method — locators as private fields in the constructor, actions as public methods.')
numbered('Write the test — beforeEach for common navigation, meaningful TC ID, tag @smoke/@regression.')
numbered('Run locally, headed, against the real target environment.')
numbered('Self-review before pushing — Claude as the reviewer there is no second engineer to provide.')
numbered('Run npx tsc --noEmit — TypeScript correctness gate before any push.')
numbered('Push and watch CI — manually triggered, choose environment, verify on Azure cloud browsers.')
numbered('If it fails in CI but passed locally — first suspect env/credentials/CI-runner differences, not test logic.')
numbered('Document the fix in Fixes.md immediately, while the root cause is fresh.')
soundbite("Working solo means no peer-review gate, so I built two things to compensate: Claude as a "
          "standing code reviewer before every push, and a structured Fixes.md so that once I've solved a "
          "class of problem -- like a dropdown race condition -- I check it against every other page "
          "object instead of only fixing the one instance in front of me.")

# ============================ 18. CODE REVIEW GUIDELINES ============================
h1('18. Code review guidelines')
para('Organised into categories with reasoning — not read back as a flat checklist.')

h2('Readability')
bullet('Self-explanatory names over comments — clickDraftDocumentCheckbox() already says what it does.')
bullet('Comments explain WHY, not WHAT — e.g. our comment on contains(@href,...) explaining WHY it is '
       'not an exact match (preprod URL prefix) is a good comment. Restating the XPath in English is not.')

h2('Consistency')
bullet('Follow agreed conventions — locators-in-constructor, methods-grouped-by-section is identical '
       'across all four page objects; a reviewer should never relearn structure per file.')
bullet('Consistent formatting/indentation — config-enforced, not eyeballed.')

h2('Maintainability')
bullet('No hardcoded values — PARTNER_CATEGORY_NAME, DOCUMENT_NAME come from .env/config.ts, because the '
       'value differs per environment.')
bullet('DRY — the expect().toPass() retry pattern is one piece of logic reused across two page objects, '
       'not copy-pasted with subtle variations.')

h2('Correctness')
bullet('Strict-mode risk — does this locator resolve to exactly ONE element on every environment, not '
       'just the one tested on? (Fix 3, Fix 4.)')
bullet('Race conditions — is there an action right after a navigation-triggering click that could run '
       'against a destroyed page context? (Fix 9.)')
bullet('Edge cases / negative paths — happy path only, or also missing required fields, special '
       'characters, wrong file sizes?')

soundbite('My checklist is not generic — every item maps to a real bug I have hit in this codebase. '
          'Strict-mode violations, silent contains() vs exact-match failures, and evaluate() running '
          'mid-navigation are things I check for specifically now because I was burned by each one once.')

# ============================ 19. STORAGESTATE CORE CONCEPTS ============================
doc.add_page_break()
h1('19. storageState — Core Concepts & Mechanics')

h2('19.1 — Purpose, and how it saves time')
para("Without it, every test repeats the full UI login flow: navigate -> fill username -> fill password "
     "-> submit -> wait for redirect. Our auth.setup.ts does that ONCE, then saves the result with "
     "page.context().storageState({ path: 'auth.json' }). Every other test loads that file into a brand-"
     "new browser context instead of repeating the login flow.")
para("Time math: if login takes ~5-8s and there are 50 tests, doing it once instead of 50 times saves "
     "roughly 4-6 minutes per run -- and removes 49 chances for the login UI itself to flake and take "
     "down an unrelated test.")

h2('19.2 — What is actually inside a generated auth.json')
code('{\n'
     '  "cookies": [\n'
     '    { "name": "...", "value": "...", "domain": "...", "path": "/", "expires": 123456,\n'
     '      "httpOnly": true, "secure": true, "sameSite": "Lax" }\n'
     '  ],\n'
     '  "origins": [\n'
     '    { "origin": "https://app.sppreprod.in",\n'
     '      "localStorage": [ { "name": "...", "value": "..." } ] }\n'
     '  ]\n'
     '}')
para('Two top-level keys, always: cookies + localStorage, per origin. Nothing else — that "nothing else" '
     'is the whole answer to 19.3.')

h2('19.3 — Why sessionStorage is excluded, and how to bypass it')
para('Not an oversight — deliberate. By spec, sessionStorage is scoped to ONE tab\'s lifetime. Even a '
     'real user opening a new tab does not inherit the old tab\'s sessionStorage. Playwright models '
     '"returning as a logged-in user," not "resuming the exact same tab," so it correctly leaves it out.')
code("// Capture (once, e.g. inside auth.setup.ts after login)\n"
     "const sessionData = await page.evaluate(() => JSON.stringify(sessionStorage));\n\n"
     "// Restore (addInitScript runs BEFORE the app's own JS, on every new document)\n"
     "await context.addInitScript(data => {\n"
     "  const parsed = JSON.parse(data);\n"
     "  for (const [key, value] of Object.entries(parsed)) {\n"
     "    window.sessionStorage.setItem(key, value as string);\n"
     "  }\n"
     "}, sessionData);")

h2('19.4 — globalSetup vs Project Dependencies')
para('We actually use BOTH, for two different jobs.')
table(['', 'globalSetup (our global-setup.ts)', 'Project Dependencies (our setup -> chromium)'],
      [['Runs as', 'A plain function, once, separate process, before anything else', "A REAL Playwright test (test as setup)"],
       ['Fixture access', 'None unless you manually chromium.launch()', 'Full — page, browser, auto-wait, everything'],
       ['Shows in reports?', 'No', 'Yes — own named step, trace/screenshot on failure'],
       ['Retries on failure?', 'No', 'Yes — it is a normal test'],
       ['Good for', 'Non-test housekeeping (our case: wiping allure-results/)', 'Anything needing auto-wait/retry/trace — login is exactly this']])
soundbite("I use globalSetup for the one thing that doesn't need a browser -- clearing the Allure results "
          "folder -- and project dependencies for login, specifically because login deserves trace and "
          "screenshot capture if it breaks. If login fails, every downstream test fails with it.")

h2('19.5 — Why playwright/.auth/ + .gitignore')
bullet('Directory convention is about SCALING, not security — one folder lets you .gitignore one glob '
       '(playwright/.auth/*) instead of a new line per file as roles are added.')
bullet('.gitignore is about security — auth.json is a LIVE, valid session. Anyone holding it can act as '
       'that logged-in user with no password, until it expires. Committing it once leaks it into git '
       'history forever, even if deleted in a later commit.')
fixme('We already gitignore auth.json correctly, but it sits at the project root instead of inside a '
      'playwright/.auth/ folder. Worth fixing before adding multi-role files (section 21.1) — one folder '
      'scales cleaner than scattering admin.json/editor.json/customer.json at root.')

# ============================ 20. STORAGESTATE CONFIG & IMPLEMENTATION ============================
doc.add_page_break()
h1('20. storageState — Configuration & Implementation')

h2('20.1 — Wiring a project to consume auth.json (this IS our config)')
code("{\n"
     "  name: 'chromium',\n"
     "  use: { ...devices['Desktop Chrome'], storageState: 'auth.json' },\n"
     "  dependencies: ['setup'],\n"
     "}")
para('use.storageState hydrates every new context in this project with cookies/localStorage from the '
     'file. dependencies: [\'setup\'] guarantees login runs first and succeeds before any test starts.')

h2('20.2 — Capturing state with context.storageState()')
code("// Write directly to a file\n"
     "await page.context().storageState({ path: 'auth.json' });\n\n"
     "// OR get it back as an object (needed for the API-combo pattern in 20.4)\n"
     "const state = await page.context().storageState();")

h2('20.3 — Bypassing global storageState for one guest/unauthenticated spec file')
code("// tests/e2e/guest-checkout.spec.ts\n"
     "import { test, expect } from '../../utils/fixtures';\n\n"
     "test.use({ storageState: { cookies: [], origins: [] } }); // empty = unauthenticated\n\n"
     "test('guest can browse without logging in', async ({ page }) => {\n"
     "  await page.goto('/home');\n"
     "});")
para('test.use() overrides the project-level config only for tests declared after it in that file — every '
     'other spec file keeps using auth.json untouched.')

h2('20.4 — Combine API-based auth with storageState (eliminate UI login entirely)')
para('Upgrade path for our own auth.setup.ts, which is currently UI-based:')
code("setup('authenticate via API', async ({ request, browser }) => {\n"
     "  const apiContext = await request.newContext({ baseURL: BASE_URL });\n\n"
     "  const response = await apiContext.post('/api/auth/login', {\n"
     "    data: { email: USER_EMAIL, password: USER_PASSWORD },\n"
     "  });\n"
     "  expect(response.ok()).toBeTruthy();\n\n"
     "  const state = await apiContext.storageState();\n"
     "  const context = await browser.newContext({ storageState: state });\n"
     "  await context.storageState({ path: 'auth.json' });\n"
     "});")
para('If the session is a Bearer token in localStorage instead of a cookie (common for SPA apps like '
     'ours), inject it via addInitScript before the page loads:')
code("const context = await browser.newContext();\n"
     "await context.addInitScript(tok => window.localStorage.setItem('authToken', tok), token);\n"
     "const page = await context.newPage();\n"
     "await page.goto('/');\n"
     "await context.storageState({ path: 'auth.json' });")
soundbite("Our login flow has waitForURL('**/AssetLibrary', { timeout: 60000 }) -- a 60-second budget for "
          "a redirect. An API login removes the entire UI render/redirect chain; setup would likely drop "
          "from several seconds to under one.")

h2('20.5 — Strategy for IndexedDB-based sessions')
para('Same bucket as sessionStorage (19.3), but heavier — IndexedDB is a full async, versioned database, '
     'not a flat key-value store. Capture every object store into a plain object, persist to JSON, then '
     'replay via addInitScript before the app\'s own code runs. Build it once as a small reusable helper '
     '(utils/indexedDbState.ts with captureIndexedDb()/restoreIndexedDb()) rather than inlining per test.')

# ============================ 21. STORAGESTATE ADVANCED & SCALING ============================
doc.add_page_break()
h1('21. storageState — Advanced Architecture & Scaling')

h2('21.1 — Multi-role authentication (admin / editor / customer)')
para('Pattern A — separate projects per role, when roles never mix inside one test:')
code("projects: [\n"
     "  { name: 'setup-admin', testDir: '.', testMatch: /admin\\.setup\\.ts/ },\n"
     "  { name: 'setup-customer', testDir: '.', testMatch: /customer\\.setup\\.ts/ },\n\n"
     "  { name: 'admin-tests', use: { storageState: 'playwright/.auth/admin.json' },\n"
     "    dependencies: ['setup-admin'], testMatch: /.*\\/admin\\/.*\\.spec\\.ts/ },\n\n"
     "  { name: 'customer-tests', use: { storageState: 'playwright/.auth/customer.json' },\n"
     "    dependencies: ['setup-customer'], testMatch: /.*\\/customer\\/.*\\.spec\\.ts/ },\n"
     "]")
para('Pattern B — explicit multi-context inside ONE test, when one test needs both roles at once:')
code("test('admin creates, customer sees it', async ({ browser }) => {\n"
     "  const adminCtx = await browser.newContext({ storageState: 'playwright/.auth/admin.json' });\n"
     "  const adminPage = await adminCtx.newPage();\n"
     "  // ...admin creates something\n\n"
     "  const customerCtx = await browser.newContext({ storageState: 'playwright/.auth/customer.json' });\n"
     "  const customerPage = await customerCtx.newPage();\n"
     "  // ...customer verifies it\n"
     "});")
para('Same idea as the Promise.all multi-tab/parent-child window pattern — just for role isolation '
     'instead of parent/child windows.')

h2('21.2 — Avoiding race conditions with multiple workers sharing one auth.json')
para("We currently avoid this entirely -- workers: 1 and fullyParallel: false are deliberate, our own "
     "config comment says exactly why: shared session, avoid conflicts.")
bullet('The file itself is not at risk — loading auth.json into N contexts is safe, each gets its own '
       'in-memory cookie jar, nobody writes back to the file.')
bullet('The REAL risk is server-side — single-session-per-user enforcement, CSRF token rotation, or '
       'session storage keyed in a way that two simultaneous logins on the same account collide. This is '
       'an application behaviour, not a Playwright limitation.')
bullet('Actual fix for true parallelism: one dedicated test account PER worker, each with its own setup '
       'run and its own state file — no two workers ever share a server-side session.')
soundbite("I haven't needed to solve this yet because I deliberately chose workers: 1 to avoid it -- but "
          "the real fix isn't a Playwright setting, it's asking the backend team whether their session "
          "model tolerates concurrent logins on one account. If it doesn't, the answer is one test account "
          "per worker, not a config trick.")

h2('21.3 — MFA / OTP in automated auth setup')
numbered('Best: get the test account MFA-exempted at the environment/feature-flag level — the standard '
         'industry answer; do not fight the security layer from the test layer if avoidable.')
numbered('TOTP (rotating code) — fully automatable. Retrieve the seed once, generate the current code '
         'deterministically with a library, no waiting:')
code("import { TOTP } from 'otpauth';\n"
     "const totp = new TOTP({ secret: process.env.MFA_SECRET });\n"
     "await page.locator('#otp').fill(totp.generate());")
numbered('Email/SMS OTP — never read a real inbox via UI. Use a programmatic test-inbox API (Mailosaur, '
         'a dedicated test IMAP account, or an internal endpoint exposing the last-sent code).')
numbered('Last resort — manual/mocked, only acceptable in a throwaway demo environment.')

h2('21.4 — Session cookie expires mid-CI-run')
para('Same CLASS of problem as our own Fix 18 (test timeout exhausted under load) — just at the '
     'session-TTL layer. Symptom: tests pass for the first N minutes, then start failing with 401s/'
     'login-redirects with no code change.')
bullet('Confirm session TTL with the backend comfortably exceeds the suite\'s total runtime.')
bullet('Shard the run (--shard=1/4) so each shard runs its own setup project independently — no single '
       'session survives the entire suite, and shards run in parallel besides.')
bullet('If the app supports refresh tokens, proactively refresh before expiry rather than reacting to a 401.')
bullet('Fallback: detect a 401 mid-suite, re-run setup login, retry once — better as a deliberate check '
       'than relying on Playwright retries to mask it.')

h2('21.5 — 401 despite explicitly loading auth.json — debugging order')
numbered('Check expires on the relevant cookie inside auth.json — already in the past?')
numbered('Check the cookie\'s domain against the baseURL actually being run against.')
fixme('This one is a REAL risk in our specific framework — four environments (dev/preprod/prod/digipulse) '
      'share ONE auth.json filename. Regenerate it under ENV=dev, then run with ENV=prod, and you get '
      'exactly this 401 — looks like a Playwright bug, is actually a stale-file-for-the-wrong-environment bug.')
numbered('Check for a CSRF/anti-forgery token sent as a custom header, not a cookie — storageState only '
         'captures cookies + localStorage, never custom headers.')
numbered('Verify project wiring — a typo\'d storageState path or a missing dependencies: [\'setup\'] '
         'silently runs with NO session, not a config error.')
numbered('Check IP/User-Agent binding — some backends invalidate a session if the request origin changes '
         '(generated locally, consumed by a CI runner/Azure cloud browser at a different IP).')
numbered('Circle back to 19.3/20.5 — if the app\'s "are you logged in" check reads sessionStorage or '
         'IndexedDB, a structurally complete auth.json loads fine and still gets rejected.')
soundbite("Given that this framework runs four environments off one auth.json filename, the domain-"
          "mismatch case is not theoretical for me -- it's the first thing I'd suspect if I ever saw a 401 "
          "right after switching ENV without regenerating the session first.")

# ============================ FOOTER ============================
doc.add_paragraph()
foot = doc.add_paragraph()
fr = foot.add_run('Built from interview-prep Q&A sessions, grounded in this repo\'s real config, fixtures, '
                  'and Fixes.md. Companion to Concepts-Study-Guide.docx (JS/TS language fundamentals) and '
                  'PROGRESS.md (what was built, session by session).')
fr.italic = True
fr.font.size = Pt(9)

doc.save('Playwright-Interview-Prep.docx')
print('OK: Playwright-Interview-Prep.docx written')
