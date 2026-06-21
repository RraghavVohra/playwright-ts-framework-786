# -*- coding: utf-8 -*-
"""Generates Concepts-Study-Guide.docx from in-script content."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- base styles ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

ACCENT = RGBColor(0x1F, 0x4E, 0x79)   # dark blue
CODEBG = 'F2F2F2'

def shade(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)

def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p

def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def numbered(text):
    return doc.add_paragraph(text, style='List Number')

def code(text):
    p = doc.add_paragraph()
    shade(p, CODEBG)
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    # ensure monospace applies for east-asian too
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), 'Consolas')
    rfonts.set(qn('w:hAnsi'), 'Consolas')
    return p

def soundbite(text):
    p = doc.add_paragraph()
    shade(p, 'FFF2CC')
    p.paragraph_format.left_indent = Inches(0.1)
    r = p.add_run('Interview soundbite:  ')
    r.bold = True
    r2 = p.add_run(text)
    r2.italic = True
    return p

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, htext in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(htext)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()
    return t

# ============================ TITLE ============================
title = doc.add_heading('JavaScript + TypeScript Concepts', level=0)
sub = doc.add_paragraph()
sr = sub.add_run('A Study Guide for Interviews  —  taught simply, compared to Java, '
                 'grounded in our Playwright + TypeScript framework')
sr.italic = True
sr.font.size = Pt(11)
doc.add_paragraph()
para('How to read this: each concept has four parts — the idea, the Java comparison, '
     'our actual code, and an interview soundbite you can say out loud.', italic=True)
doc.add_page_break()

# ============================ 0 ============================
h1('0. The Big Picture: JS + TS')
para('JavaScript (JS) is the language that actually runs. TypeScript (TS) is a layer on top '
     'that adds types — labels saying "this value must be a number / string / Page". The browser '
     'and Node do not run TS directly; it is compiled to JS first.')
bullet('JS gives us: how the code behaves — functions, objects, Promises, modules.')
bullet('TS adds: safety — it catches type mistakes before the code runs.')
para('Mental model: JS is the engine. TS is the safety inspector checking your parts fit before '
     'you start it. Coming from Java (also strongly typed), TS gives back the type safety you know, '
     'on top of JavaScript.', italic=True)

# ============================ 1 FUNCTIONS ============================
h1('1. Functions')
h2('The idea')
para('A function is a reusable block of code that takes inputs and optionally returns an output. '
     'The biggest shift from Java: in JavaScript a function is a value — it can be stored in a '
     'variable, passed as an argument, or returned from another function. This is called '
     '"functions are first-class citizens".')
h2('The Java comparison')
para('In Java a function must live inside a class (a method) — it cannot exist on its own. In JS a '
     'function can float freely and be passed around as data. That single difference explains '
     'fixtures, callbacks, and much of what looks "weird" coming from Java.')
h2('The four forms we use in this project')
table(['Form', 'Looks like', 'Where in our project', 'Java analogue'],
      [['Function declaration', 'function globalSetup() {}', 'utils/global-setup.ts', 'a static method'],
       ['Class method', 'async navigateToTestimonials() {}', 'every page object', 'instance method'],
       ['Arrow function (as a value)', 'async ({ page }, use) => {}', 'utils/fixtures.ts', 'lambda (closest)'],
       ['Callback (passed as arg)', 'evaluate(el => el.click())', 'page objects', 'lambda + functional interface']])
h2('Typing a function (the TS layer)')
code("getFutureDate(daysFromToday: number): string { ... }\n"
     "//            \\____ param ____/        \\ return type")
bullet('Parameter: name: type — the type comes AFTER the name (Java is reverse: int days).')
bullet('Return type: : string after the ().')
bullet("Async functions return Promise<void> — the async version of Java's void.")
soundbite('"I use three forms depending on context — function declarations for hooks, class methods '
          'for Page Object actions, and arrow functions wherever a function is used as a value, like '
          'fixtures and callbacks. The big shift from Java is that functions are first-class values "'
          'in JS. TypeScript only adds the type layer on top."')

# ============================ 1B FUNCTIONS (PLAIN) ============================
doc.add_page_break()
h1('1B. Functions — Plain & Simple (Revision)')
para('Same topic as section 1, written in the simplest language for quick revision.', italic=True)

h2('1. What a function is')
para('A function is a named block of code that does a job. You give it inputs, it runs, and it can '
     'give back a result. You write it once and call it as many times as you need — the point is to '
     'avoid repeating code.')

h2('2. How you write one')
code("function add(a: number, b: number): number {\n  return a + b;\n}")
para('Piece by piece:')
bullet('function — the keyword that starts a function.')
bullet('add — the name you give it.')
bullet('(a: number, b: number) — the inputs (parameters). Each is written as name: type.')
bullet('": number" — the return type: this function gives back a number.')
bullet('return a + b; — the result it hands back.')
para('You call it like this:')
code("const sum = add(2, 3);   // sum is now 5")

h2('3. Java vs TypeScript — same idea, two differences')
code("// Java\nint add(int a, int b) {\n  return a + b;\n}\n\n"
     "// TypeScript\nfunction add(a: number, b: number): number {\n  return a + b;\n}")
numbered('The type comes AFTER the name. Java: int a. TypeScript: a: number. Same info, reversed order.')
numbered('The return type goes at the end, after the (), not at the front.')
para('Everything else is the same: parameters, return, calling it.')

h2('4. The one new idea: a function is a value')
para('In Java a function must live inside a class — it cannot exist on its own. In JavaScript a '
     'function is a value, the same way 5 is a value or "hello" is a value. So you can store a '
     'function in a variable, pass it into another function, or return it from a function.')
para('Example — storing a function in a variable:')
code('const greet = function(name: string): string {\n  return "Hello " + name;\n};\n\n'
     'greet("Raghav");   // "Hello Raghav"')
para('Here greet is a variable that holds a function. This single idea is what makes fixtures and '
     'callbacks work later.')

h2('5. The arrow function (a shorter way to write a function)')
para('JavaScript has a shorter syntax using =>, called an arrow function. These two are the same:')
code("// normal function\nconst add = function(a: number, b: number): number {\n  return a + b;\n};\n\n"
     "// arrow function — same thing, shorter\nconst add = (a: number, b: number): number => {\n  return a + b;\n};")
para('The arrow function drops the word function and puts => after the parameters. You will see '
     'arrow functions everywhere in Playwright.')

h2('6. async functions (a short note for now)')
para('Many functions in our project start with async:')
code("async navigateToTestimonials(): Promise<void> {\n  await this.setupTab.click();\n}")
bullet('async means the function does work that takes time (like clicking in a browser).')
bullet("Promise<void> is its return type — the 'takes time, returns nothing' version of Java's void.")
para('We cover async, await, and Promise fully in their own section. Do not worry about them here.')

h2('7. The forms of functions we use in this project')
para('a) A standalone function — utils/global-setup.ts (closest to a Java method):')
code("function globalSetup() {\n  // delete old report files before tests run\n}")
para('b) A method inside a class — every page object (inside a class you do NOT write the word function):')
code("async navigateToTestimonials(): Promise<void> {\n  await this.setupTab.click();\n  await this.testimonialsNewOption.click();\n}")
para('c) A function used as a value — utils/fixtures.ts (the "function is a value" idea in real use; '
     'see the deep dive below):')
code("testimonialsPage: async ({ page }, use) => {\n  const testimonialsPage = new TestimonialsPage(page);\n  await use(testimonialsPage);\n}")

h2('8. Deep dive: a function as the value of an object key')
para('Point 7c confuses people, so here it is step by step.')
para('Step 1 — a function can be a value (from point 4): const greet = function() {...}. greet is a '
     'variable whose value is a function.')
para('Step 2 — an object is a collection of key: value pairs:')
code('const person = {\n  name: "Raghav",\n  age: 25,\n};')
para('name and age are keys; "Raghav" and 25 are their values.')
para('Step 3 — since a function is a value, a value in an object can be a function:')
code('const person = {\n  name: "Raghav",\n  greet: function() {\n    return "Hello";\n  },\n};')
para('Now the key greet has a function as its value. That is the whole trick.')
para('Step 4 — our fixtures file is one big object passed into base.extend(...). Each key is a '
     'fixture, each value is a function:')
code("base.extend({\n  testimonialsPage: async ({ page }, use) => {\n"
     "    const testimonialsPage = new TestimonialsPage(page);\n    await use(testimonialsPage);\n  },\n});")
para('Important: base.extend(...) is a FUNCTION CALL. The { ... } you pass into it is the object. '
     'That object has keys (fixture names) whose values are functions. So the structure is simply: '
     'an object key whose value is a function.')

h2('The one line to remember')
para('In Java a function must live inside a class. In JavaScript a function is a value — you can '
     'store it, pass it, and return it. TypeScript only adds the types (name: type for inputs, '
     ': type for the return).', italic=True)

# ============================ 2 MODULES ============================
h1('2. Modules: import / export')
h2('The idea')
para('In JavaScript every file is a module. A file is a sealed box: nothing inside it is visible to '
     'other files unless you export it. Other files then import what they need.')
h2('The Java comparison')
table(['Java', 'TypeScript'],
      [['Organise by classes & packages', 'Organise by files'],
       ['import com.app.Config; then Config.ENV', "import { ENV } from './utils/config' — value directly"],
       ['package-private (no modifier)', 'no export = file-private'],
       ['Import a class, then dot into it', 'The file itself is the namespace']])
h2('Named exports — a toolbox of several things')
code("export const ENV = ...;        // exported — importable\n"
     "export const BASE_URL = ...;   // exported\n"
     "const BASE_URLS = { ... };     // NO export — private to this file")
para('Import named exports with curly braces, using the exact name:')
code("import { ENV, USER_EMAIL } from './utils/config';")
para('Key encapsulation point: BASE_URLS (the full map) has no export, so it is file-private — like '
     'a private field. Only BASE_URL (the one resolved URL) is exported — like a public getter. '
     'Other files get the result, never the internal map.')
h2('Default exports — the one main thing')
code("export default async function globalSetup() { ... }\n"
     "// imported WITHOUT braces, and you choose the name:\n"
     "import globalSetup from './utils/global-setup';")
h2('Re-exporting & aliasing — our fixtures file')
code("import { test as base, expect } from '@playwright/test';  // 'as' = rename on import\n"
     "export const test = base.extend(...);                     // build our own 'test'\n"
     "export { expect };                                        // re-export\n"
     "// result: tests need ONE import:\n"
     "import { test, expect } from '../../utils/fixtures';")
h2('Relative paths')
bullet('./ = same folder · ../ = one folder up · counted from the file you are in.')
bullet('A bare name (@playwright/test, fs, path) = a package or Node built-in.')
bullet('No .ts extension in the path — TypeScript resolves it.')
bullet('If you move a file, every ../ count changes — a misplaced spec and a broken import appear together.')
soundbite('"Every file is a module. Named exports are imported with braces by exact name; anything '
          'without export stays file-private — encapsulation at the file level. I use a default '
          'export for global-setup, and in fixtures I alias and re-export so every test has one clean import line."')

# ============================ 2B MODULES (PLAIN) ============================
doc.add_page_break()
h1('2B. Modules (import / export) — Plain & Simple (Revision)')
para('Same topic as section 2, written in the simplest language for quick revision.', italic=True)

h2('1. What a module is')
para('In JavaScript, every file is a module. One file = one module. A module is like a sealed box: '
     'by default, whatever you write inside a file stays inside that file. No other file can see it.')

h2('2. The problem modules solve')
para('A project has many files (config.ts, fixtures.ts, TestimonialsPage.ts ...) and they need to '
     'share code. config.ts has the URL that test files need; TestimonialsPage.ts has the class that '
     'fixtures.ts needs. So we need a way to send code out of a file (export) and bring code into a '
     'file (import). That is all modules are: a system for files to share code.')

h2('3. export — open a window in the box')
para('export means: make this thing visible to other files.')
code('export const ENV = "preprod";          // other files CAN use this\n'
     'export const BASE_URL = "https://...";  // other files CAN use this\n\n'
     'const BASE_URLS = { ... };              // NO export -> trapped inside this file')
para('The first two can be used by other files. BASE_URLS has no export, so it is private to this '
     'file — no other file can touch it.')

h2('4. import — take something from another file')
para('import means: bring in a thing another file exported. From auth.setup.ts:')
code("import { ENV, USER_EMAIL, USER_PASSWORD } from './utils/config';")
para('This reads as: go to the file ./utils/config and bring in ENV, USER_EMAIL, USER_PASSWORD. '
     'config.ts exports about 15 things, but this file imports only the 3 it needs. You take only '
     'what you want, not the whole file.')

h2('5. Java comparison')
table(['Java', 'TypeScript'],
      [['Code grouped in classes and packages', 'Code grouped in files'],
       ['import com.app.Config; then Config.ENV', "import { ENV } from './utils/config'; then ENV directly"],
       ['no modifier = package-private', 'no export = file-private'],
       ['Import a class, then dot into it', 'Import the value itself — no class wrapper']])
para('Biggest switch from Java: there is no class wrapping these values. In Java you do Config.ENV. '
     'In TypeScript the file itself is the container, so you import ENV and use it directly.')

h2('6. Two kinds of export')
para('a) Named export — export things by their name. A file can have many.')
code('// config.ts\nexport const ENV = "preprod";\nexport const BASE_URL = "https://...";')
para('Import them WITH curly braces, using the exact same name:')
code("import { ENV, BASE_URL } from './utils/config';")
para('b) Default export — one file, one main thing. No braces.')
code('// global-setup.ts\nexport default function globalSetup() { ... }')
para('Import it WITHOUT braces, and you may call it any name:')
code("import globalSetup from './utils/global-setup';")
para('Quick rule: many things to share -> named exports (braces). One main thing -> default export (no braces).')

h2('7. The path part — from \'./utils/config\'')
para('The text after from tells JavaScript where to find the file:')
bullet('./ = the same folder.')
bullet('../ = go up one folder.')
bullet("A plain name like '@playwright/test' or 'fs' = an installed package or built-in tool, not your own file.")
para('Example from a test file two folders deep:')
code("import { test, expect } from '../../utils/fixtures';")
para('../../ means go up two folders to the project root, then into utils. Important: the path is '
     'counted FROM the file you are writing in. That is why a file in the wrong folder also has the '
     'wrong ../ count — the path depends on where the file physically sits.')

h2('The one line to remember')
para('Every file is a sealed box. export makes something visible to other files; import pulls it in. '
     'Named exports use braces and the exact name; a default export uses no braces and any name. '
     'Unlike Java, there is no class wrapper — the file itself is the container.', italic=True)

# ============================ 3 PROMISES ============================
h1('3. Promises & async/await')
h2('The idea')
para('A Promise is an object representing a value that is not ready yet but will be — an IOU. It is '
     'always pending, fulfilled (success + value), or rejected (failure + error).')
h2('Why JS needs Promises (the crucial fact)')
para('JavaScript runs on a single thread. It cannot afford to block and wait, because freezing the '
     'one thread freezes everything. So slow operations (clicks, navigation, network) immediately '
     'return a Promise instead of blocking.')
h2('The Java comparison')
bullet('A Promise is like Java CompletableFuture<T> — a container for a result that arrives later.')
bullet('await somePromise is like future.get() — wait for the result and unwrap it.')
bullet('Difference: Java .get() blocks the whole thread. JS await only pauses the current function — '
       'the single thread stays free. This is non-blocking concurrency.')
h2('async and await')
code("await testimonialsPage.navigateToTestimonials();   // wait for completion (Promise<void>)\n\n"
     "const heading = await page.getPageHeading();        // wait AND unwrap the value\n"
     "expect(heading).toBe('PUSH NOTIFICATION');          // heading is now a real string")
h2('The #1 bug: forgetting await')
code("const heading = page.getPageHeading();   // no await\n"
     "expect(heading).toBe('PUSH NOTIFICATION'); // compares a Promise to a string -> FAILS")
para('Without await, heading holds the Promise object itself, not the string — a wrong TYPE. With '
     'actions, forgetting await makes the next line race ahead before the action finishes — the #1 '
     'cause of flaky tests.')
h2('Which expect needs await?')
code("expect(heading).toBe('PUSH NOTIFICATION');       // NO await — value already in hand (sync)\n"
     "await expect(page).toHaveURL(/testimonial/);     // await — polls the LIVE page, retries (async)")
para('Rule: await expect(...) when the matcher checks the live page (toHaveURL, toBeVisible, '
     'toHaveText). Plain expect(...) when asserting on a value you already have (toBe, toEqual).')
soundbite('"JavaScript is single-threaded, so it cannot block while waiting for the browser. Slow '
          'operations return a Promise — like Java CompletableFuture. async makes a function return a '
          'Promise; await pauses that function until it resolves and unwraps the value. Forgetting '
          'await is the top cause of flaky tests. Unlike Java .get(), await only pauses the current '
          'function, not the whole thread."')

# ============================ 3B PROMISES (PLAIN) ============================
doc.add_page_break()
h1('3B. Promises & async/await — Plain & Simple (Revision)')
para('Same topic as section 3, written in the simplest language for quick revision.', italic=True)

h2('1. Start with the problem (the why)')
para('Everything Playwright does takes time: clicking a button, loading a page, reading text. The '
     'question is: what does your code do while it waits? To answer that, you need one fact: '
     'JavaScript runs on a single thread — it has only ONE worker.')
para('If JavaScript stopped and waited every time it clicked, that one worker would freeze and '
     'nothing else could happen — the whole program would lock up. So it cannot afford to wait. '
     'That is why it needs Promises.')

h2('2. What a Promise is')
para('A Promise is an object that represents a result that is not ready yet, but will be. Think of '
     'it like a receipt at a restaurant: you order food, you do not get it instantly — you get a '
     'receipt (the Promise) that says "your food is coming". Later the food arrives (fulfilled) or '
     'the kitchen runs out (rejected).')
para('A Promise is always in one of three states:')
bullet('pending — still working, no result yet.')
bullet('fulfilled — finished, here is the result.')
bullet('rejected — failed, here is the error.')
para('So page.click() does not freeze — it hands back a Promise that says "the click is happening, '
     'I will tell you when it is done".')

h2('3. async — what / how / why')
para('What: a keyword you put in front of a function. How:')
code("async navigateToTestimonials(): Promise<void> {\n  ...\n}")
para('Why: putting async on a function makes it return a Promise automatically. It tells JavaScript '
     '"this function does slow work — give callers a receipt, not an instant answer". That is why the '
     'return type is Promise<void>: Promise = a receipt; <void> = no value when it finishes, it just '
     'completed.')

h2('4. await — what / how / why')
para('What: a keyword you put in front of a Promise. How:')
code("await this.setupTab.click();")
para('Why: await means "pause here, wait for this Promise to finish, then continue". It does two '
     'jobs: (1) waits for the slow work, and (2) unwraps the result — takes the value out of the '
     'Promise and gives it to you.')
code("await this.setupTab.click();              // job 1: wait until the click is done\n\n"
     "const heading = await page.getPageHeading();  // job 1 + 2: wait AND take the value out")
para('In the second line getPageHeading() returns a Promise containing a string; await pulls the '
     'string out, so heading is now a normal string.')

h2('5. The key point about await and the single thread')
para('When JavaScript hits await, it pauses only that one function — but the single worker is set '
     'free to do other things meanwhile. When the Promise is ready, it comes back and continues. '
     'This is the big difference from Java:')
table(['Java', 'JavaScript'],
      [['future.get() blocks the whole worker', 'await pauses only the current function'],
       ['other work waits', 'the one worker stays free for other work']])
para('This is called non-blocking. It is why JavaScript can wait for slow things without freezing.')

h2('6. The #1 mistake: forgetting await')
code("const heading = page.getPageHeading();   // forgot await\n"
     "expect(heading).toBe('PUSH NOTIFICATION');")
para('Without await, heading is the Promise (the receipt), not the string — so you compare a receipt '
     'to a string and it fails. With an action, forgetting await means the next line runs immediately, '
     'before the action finished. This is the number one cause of flaky tests. That is why every '
     'browser action in our framework has await in front of it.')

h2('7. Which expect needs await?')
code("expect(heading).toBe('PUSH NOTIFICATION');      // NO await\n"
     "await expect(page).toHaveURL(/testimonial/);    // await")
bullet('expect(heading).toBe(...) — heading is already a string (unwrapped earlier). Instant check, no await.')
bullet('await expect(page).toHaveURL(...) — keeps re-checking the LIVE page until it matches. Takes time, returns a Promise, needs await.')
para('Rule: await expect(...) when it checks the live page (toHaveURL, toBeVisible, toHaveText). '
     'Plain expect(...) when you check a value you already have (toBe, toEqual on a string/number).')

h2('8. async vs Promise vs await — comparison table')
table(['', 'async', 'Promise', 'await'],
      [['What it is', 'a keyword on a function', 'an object (a receipt)', 'a keyword before a Promise'],
       ['Where it goes', 'in front of a function', 'returned by an async function', 'in front of a Promise/async call'],
       ['Its job', 'makes the function return a Promise', 'holds a result coming later', 'pauses until ready, then unwraps the value'],
       ['Java analogue', 'method returning CompletableFuture', 'CompletableFuture<T>', 'future.get() (but non-blocking)'],
       ['In our code', 'async navigateToTestimonials()', 'Promise<void> return type', 'await this.setupTab.click()']])
para('Link between them: async makes a function hand back a Promise, and await opens that Promise to '
     'get the value.')

h2('9. Visual')
try:
    doc.add_picture('async-await-flow.png', width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    para('[diagram async-await-flow.png not found]', italic=True)

h2('The one line to remember')
para('JavaScript has one worker, so it cannot freeze while waiting. Slow work returns a Promise (a '
     'receipt for a result coming later). async makes a function return a Promise. await pauses the '
     'function until the Promise is ready and unwraps the value. Forgetting await makes the next line '
     'run too early — the top cause of flaky tests.', italic=True)

# ============================ 4 OOP ============================
h1('4. OOP: Classes, Objects, Constructors')
h2('The vocabulary (identical to Java)')
bullet('Class = a blueprint. (TestimonialsPage)')
bullet('Object / instance = one real thing built from the blueprint. (testimonialsPage)')
bullet('new = builds an instance and triggers the constructor.')
bullet('Field = an instance variable. (private page: Page)')
bullet('this = "this specific instance I am inside".')
h2('Our real class')
code("import { Page, Locator } from '@playwright/test';\n\n"
     "export class TestimonialsPage {\n\n"
     "  private page: Page;                    // field: name 'page', type 'Page', private\n"
     "  private setupTab: Locator;\n"
     "  private testimonialsNewOption: Locator;\n\n"
     "  constructor(page: Page) {              // TS ctor is ALWAYS the word 'constructor'\n"
     "    this.page = page;                    // store the parameter into the object's field\n"
     "    this.setupTab = page.getByText('Setup', { exact: true });\n"
     "    this.testimonialsNewOption = page.getByRole('link', { name: 'Testimonials New' });\n"
     "  }\n\n"
     "  async navigateToTestimonials(): Promise<void> {\n"
     "    await this.setupTab.click();\n"
     "    await this.testimonialsNewOption.click();\n"
     "  }\n"
     "}")
h2('Java differences to remember')
table(['Concept', 'Java', 'TypeScript'],
      [['Constructor name', 'same as class: TestimonialsPage()', "always the word 'constructor'"],
       ['Field declaration', 'private Page page; (type first)', 'private page: Page; (name first)'],
       ['this', 'current instance', 'identical'],
       ['private', 'only this class', 'identical']])
h2('name : Type — the confusion killer')
code("private page: Page;\n"
     "//      ^      ^\n"
     "//    NAME   TYPE")
bullet('page (lowercase) = the name you chose — could be tab, browserTab, anything.')
bullet('Page (capital) = the type — the rule for what is allowed in it.')
bullet('They only look alike because we named the field after its type. Rename to tab and it is clear: private tab: Page;')
h2('The constructor — what it is REALLY for')
para('The constructor turns a freshly-created empty object into a usable one. It runs once, '
     'automatically, the instant you write new. Flow:')
numbered('new TestimonialsPage(page) creates a blank object.')
numbered('JS immediately calls the constructor, passing in page.')
numbered('this.page = page stamps the tab onto the object; locators get registered.')
para('Analogy: the page object is a TV remote; the browser tab is the TV. new TestimonialsPage(page) '
     'pairs the remote to one specific TV — and the constructor is the pairing step. Every method '
     'afterwards is a button that works because you paired first.', italic=True)
para('Without it: this.page is undefined, and the first this.page.click() crashes with "cannot read '
     'property click of undefined". The constructor is what makes this.page exist. Locators live in '
     'the constructor because it runs once — set up once, use everywhere.')
soundbite('"The constructor wires the page object to a live browser tab and registers all locators '
          'once at creation. Without it, this.page is undefined and every action crashes — the '
          'constructor turns an empty object into a usable one bound to a specific tab."')

# ============================ 4B OOP (PLAIN) ============================
doc.add_page_break()
h1('4B. OOP: Classes, Objects, Constructor — Plain & Simple (Revision)')
para('Same topic as section 4, written in the simplest language. Analogy used throughout: a '
     'blueprint and a house.', italic=True)

h2('1. Class vs Object — blueprint vs house')
para('A class is a blueprint. A blueprint for a house is not a house — it is the plan for building '
     'houses. An object is an actual house built from that blueprint. From one blueprint you can '
     'build many houses.')
code("export class TestimonialsPage { ... }                  // the blueprint\n\n"
     "const testimonialsPage = new TestimonialsPage(page);   // an actual house built from it")
para('TestimonialsPage (class) = blueprint. testimonialsPage (variable) = a real house. Same as Java.')

h2('2. Fields — what the object stores')
para('A field is information the object holds — like a house having an address. TestimonialsPage '
     'holds three things:')
code("private page: Page;                     // which browser tab this page controls\n"
     "private setupTab: Locator;              // where the Setup menu is\n"
     "private testimonialsNewOption: Locator; // where the Testimonials link is")
para('These are the object\'s belongings. Same as Java instance variables; only the syntax differs '
     '(name: type instead of type name).')

h2('3. The constructor — the setup step')
para('What: a special function that runs once, automatically, the moment you build the object with '
     'new. Its job: set up the empty object so it is ready to use.')
para('Analogy: a TV remote and a TV. The page object is the remote; the browser tab (page) is the '
     'TV. new TestimonialsPage(page) pairs the remote to one specific TV — the constructor is the '
     'pairing step. After pairing, every button (method) works because you paired first.')
code("constructor(page: Page) {\n"
     "  this.page = page;                                         // pair the remote to this TV\n"
     "  this.setupTab = page.getByText('Setup', { exact: true });\n"
     "  this.testimonialsNewOption = page.getByRole('link', { name: 'Testimonials New' });\n"
     "}")
para('Why it matters: without a constructor, this.page would be empty (undefined). The first '
     'this.page.click() would crash — "cannot read property of undefined". The constructor makes the '
     'object usable instead of empty.')
para('Java difference: in Java the constructor has the same name as the class. In TypeScript it is '
     'always the literal word constructor. That is the only syntax change.')

h2('4. this — "this particular object"')
para('this means "the specific object I am working on right now". Build two houses from one '
     'blueprint — each has its own address. this.page = page means "store the tab into THIS object\'s '
     'page field", so each object remembers its own tab.')
para('Mostly same as Java. Flag: in JavaScript this can change meaning depending on HOW a function '
     'is called. Inside a class constructor or method, this = the object (like Java). The exception '
     'comes later with arrow functions.')

h2('5. Methods — what the object can do')
para('A method is an action the object can perform.')
code("async navigateToTestimonials(): Promise<void> {\n"
     "  await this.setupTab.click();\n"
     "  await this.testimonialsNewOption.click();\n"
     "}")
para('The method uses this.setupTab and this.testimonialsNewOption — the fields the constructor set '
     'up. Setup happens in the constructor; actions happen in methods. Inside a class you do NOT '
     'write the word function.')

h2('6. new — build a house from the blueprint')
code("const testimonialsPage = new TestimonialsPage(page);")
para('What new does: (1) creates a fresh empty object, (2) runs the constructor on it (passing page), '
     '(3) hands back the finished object. Identical to Java.')

h2('7. private — who is allowed to touch it')
para('private means only code inside this class can use that field. The test cannot reach '
     'testimonialsPage.page directly — it can only call the methods the class offers. This is '
     'encapsulation: hide the inside, expose safe actions. Analogy: you use a TV through its buttons '
     '(methods), not by opening it and touching the circuit board (private fields).')
para('Java difference: in Java private is enforced at runtime. In TypeScript private is checked by '
     'the compiler only — it stops you while writing code. Behaves the same for your purposes: do '
     'not touch private fields from outside.')

h2('8. The three "page"s (the whole lesson)')
para('The word page appears three times and means three different things:')
table(['#', 'Where', 'What it is'],
      [['1', 'new TestimonialsPage(page) in the test', 'the REAL browser tab from Playwright'],
       ['2', 'constructor(page: Page)', 'a TEMPORARY parameter that receives that tab'],
       ['3', 'this.page', 'the object\'s PERMANENT field storing the tab']])
para('this.page = page is just #2 being copied into #3.')

h2('9. Visual — the journey of page through the constructor')
try:
    doc.add_picture('constructor-flow.png', width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception:
    para('[diagram constructor-flow.png not found]', italic=True)

h2('The one line to remember')
para('A class is a blueprint; an object is a real thing built from it with new. The constructor '
     'runs once at build time to set up the object\'s fields. this means "this object". Methods are '
     'what the object can do. private hides the insides so outsiders use methods, not fields.', italic=True)

# ============================ 5 LOCATORS LAZY ============================
h1('5. Locators are Lazy')
h2('The idea')
para("Writing page.getByRole('link', { name: 'Setup' }) does NOT search the page. It only describes "
     "how to find the element — like saving a recipe, not cooking. The real search happens only when "
     "you call an action on it (.click(), .textContent()).")
h2('Why this matters')
para('It is why we can safely define all locators in the constructor before navigating — defining is '
     'free, searching is deferred to action time.')
numbered('Constructor runs -> locators described (page may be blank). No searching.')
numbered('Test navigates -> real page loads.')
numbered('Method calls .click() -> NOW Playwright searches the loaded page and acts.')
h2('The Selenium contrast (great interview point)')
para('Selenium driver.findElement searches the DOM immediately — call it too early and it throws '
     'NoSuchElementException. Playwright Locator is lazy AND auto-waits — far less flaky.')
h2('Locator strategy (priority order)')
table(['Priority', 'Strategy', 'Example', 'Why'],
      [['1', 'Role + name', "getByRole('link', { name: 'Setup' })", 'resilient, user-facing — preferred'],
       ['2', 'Text', "getByText('Setup', { exact: true })", 'simple, but can match siblings'],
       ['3', 'ID', '#setup-menu', 'stable & fast when present'],
       ['4', 'XPath by text', "//a[normalize-space()='Setup']", 'when no ID/role works'],
       ['last', 'CSS class', '.nav-setup', 'classes change often — least stable']])
soundbite('"Playwright locators are lazy — defining one does not touch the DOM; it describes how to '
          'find the element. The search happens at action time, with auto-waiting built in. That is '
          'why I register every locator in the constructor before the page exists, and a big reason '
          'Playwright is less flaky than Selenium."')

# ============================ 6 POM ============================
h1('6. Page Object Model (POM)')
h2('The idea')
para('POM is a design pattern where each screen of the app gets its own class holding the locators '
     '(the elements) and the methods (the actions) for that screen. Tests call high-level methods '
     'instead of dealing with raw selectors. It is OOP applied to test automation.')
h2('Why it is good (interview gold)')
bullet('No duplication — a locator is defined once; every test reuses it.')
bullet('Easy maintenance — UI changes -> fix the locator in ONE place, not in 50 tests.')
bullet('Readable tests — the test reads like a user story, not CSS selectors.')
bullet('Encapsulation — locators are private; tests use methods, not raw selectors.')
soundbite('"POM is OOP applied to automation — one class per screen holding that screen\'s locators '
          'and actions. The payoff is maintenance: when the UI changes I update one locator in one '
          'page object, and every test that uses it is fixed."')

# ============================ 6B POM (PLAIN) ============================
doc.add_page_break()
h1('6B. Page Object Model (POM) — Plain & Simple (Revision)')
para('Same topic as section 6, written in the simplest language for quick revision.', italic=True)

h2('1. What it is')
para('POM is a simple rule: each screen of the app gets its own class. That class holds two things '
     'for that screen — the locators (where the elements are) and the methods (the actions a user can '
     'do there). TestimonialsPage is a page object: the Testimonials screen -> one class -> its '
     'locators + its navigateToTestimonials() action.')

h2('2. Why it exists')
para('Without POM, every test writes raw selectors directly:')
code("// without POM — selectors scattered across 50 tests\n"
     "await page.getByText('Setup', { exact: true }).click();\n"
     "await page.getByRole('link', { name: 'Testimonials New' }).click();")
para('If the app renames Setup to Settings, you must fix that selector in every test that used it. '
     'With POM the selector lives in one place — the page object — and the test stays clean:')
code("await testimonialsPage.navigateToTestimonials();")
para('App changes? Fix the locator once in TestimonialsPage.ts; every test is instantly fixed. That '
     'is the whole point of POM: one place to change when the UI changes.')

h2('3. The four benefits')
table(['Benefit', 'What it means'],
      [['No duplication', 'a locator is written once, reused everywhere'],
       ['Easy maintenance', 'UI change -> fix one locator, not 50 tests'],
       ['Readable tests', 'tests read like user actions, not CSS selectors'],
       ['Encapsulation', 'locators are private; tests use methods, not raw selectors']])

h2('4. POM is just OOP applied')
para('POM is not new — it is the OOP you already know, pointed at automation:')
bullet('class -> one screen')
bullet("fields -> that screen's locators")
bullet("methods -> that screen's actions")
bullet('private -> hide locators, expose actions (encapsulation)')
bullet('object -> new TestimonialsPage(page) = one live screen to drive')

h2('5. The structure it creates')
code("pages/          <- all page objects (one per screen)\n"
     "  TestimonialsPage.ts\n"
     "  PushNotificationPage.ts\n"
     "tests/e2e/      <- the tests (call page object methods)\n"
     "  testimonials.spec.ts")
para('The split is deliberate: page object = HOW to interact with a screen; test = WHAT to verify. '
     'A test never contains a selector; a page object never contains an assertion.')
table(['Thing', 'Belongs in', 'Why'],
      [['Selector / locator', 'the page object', 'one place; the test never sees raw selectors'],
       ['Assertion (expect)', 'the test', 'the test decides what to verify; the page object only knows how to interact']])

h2('The one line to remember')
para('POM means one class per screen, holding that screen\'s locators (private fields) and actions '
     '(methods). Tests call the methods instead of using raw selectors. It is OOP applied to '
     'automation — the payoff is maintenance: when the UI changes, you fix one locator in one place.', italic=True)

# ============================ 7 FIXTURES ============================
h1('7. Fixtures & Dependency Injection')
h2('The idea')
para('A fixture lets a test receive a ready-built object just by declaring it as a parameter — no '
     'manual new. You register it once; every test can ask for it.')
code("// WITHOUT a fixture (repetitive):\n"
     "test('x', async ({ page }) => {\n"
     "  const testimonialsPage = new TestimonialsPage(page);  // every test repeats this\n"
     "});\n\n"
     "// WITH a fixture (clean):\n"
     "test('x', async ({ testimonialsPage }) => {             // just ask — it is built for you\n"
     "  await testimonialsPage.navigateToTestimonials();\n"
     "});")
h2('This is Dependency Injection (DI)')
para('The test declares what it needs, and the framework provides (injects) it. Java comparison: '
     'this is exactly Spring @Autowired — declare a dependency, the framework supplies the instance.')
h2('The fixture shape (utils/fixtures.ts)')
code("testimonialsPage: async ({ page }, use) => {\n"
     "  const testimonialsPage = new TestimonialsPage(page);  // 1. SETUP — build once\n"
     "  await use(testimonialsPage);                          // 2. hand to test; test RUNS here\n"
     "  // 3. TEARDOWN — cleanup would go here (none needed yet)\n"
     "},")
para('Three edits register a new page object as a fixture: (1) import it, (2) add it to the '
     'MyFixtures type, (3) add the fixture function to base.extend({...}).')
h2('use — the most subtle part')
para('await use(obj) is NOT "return the value". It means: hand the object to the test, pause the '
     'fixture here, let the whole test run, then come back for teardown. Code before use = setup; '
     'code after use = teardown. A plain return would kill the fixture immediately — losing teardown.')
para('Analogy: use is a librarian lending you a book — fetches it (setup), waits at the desk while '
     'you read (the test runs), then shelves it when you return (teardown). return would be throwing '
     'the book and walking away.', italic=True)
soundbite('"Fixtures are Playwright\'s dependency injection — register a page object once, tests '
          'receive a ready-built instance by naming it as a parameter, same idea as Spring '
          '@Autowired. The use callback hands the object to the test and pauses the fixture around '
          'it, so code before use is setup and code after is teardown."')

# ============================ 7B FIXTURES (PLAIN) ============================
doc.add_page_break()
h1('7B. Fixtures & Dependency Injection — Plain & Simple (Revision)')
para('Same topic as section 7, written in the simplest language. Key idea: a fixture is a RECIPE '
     'you give Playwright so it can build something for you by name.', italic=True)

h2('1. The problem')
para('To use a page object you must build it with new: const testimonialsPage = new '
     'TestimonialsPage(page);. Without fixtures, every test repeats that line. 50 tests = that line '
     '50 times. We want tests to just ask for a ready-made object and get it.')

h2('2. The thing you already accept without questioning')
para('You write ({ page }) all the time, but you never write const page = new Page(). Why? Because '
     'Playwright has a built-in RECIPE for page. When you write { page }, Playwright thinks "they '
     'want page — I know how to build a browser tab", builds it, and hands it to you. You accept '
     'that magic without questioning it.')

h2('3. A fixture is you adding your OWN recipe')
para('Playwright knows page out of the box. It does NOT know testimonialsPage — that is your class. '
     'So you teach it, by registering a fixture:')
code("testimonialsPage: async ({ page }, use) => {\n"
     "  const testimonialsPage = new TestimonialsPage(page);   // the 'new' lives HERE\n"
     "  await use(testimonialsPage);\n"
     "},")
para('This hands Playwright a recipe: "if anyone asks for testimonialsPage, build it like this".')

h2('4. Why the magic is not magic')
para('When a test says async ({ testimonialsPage }) => {...}, Playwright thinks "they want '
     'testimonialsPage. Do I have a recipe for that name? Yes — they registered one. Run it." It '
     'runs your recipe (which does the new) and hands you the finished object. That is why you never '
     'wrote new in the test: the new lives inside the recipe, and Playwright runs the recipe for you '
     '— exactly like it runs its own built-in recipe for page.')

h2('5. The key test of understanding')
para('Why does { testimonialsPage } work but { bananaPage } fail? Because testimonialsPage has a '
     'registered recipe (fixture) and bananaPage does not. Playwright matches the parameter name to '
     'a registered fixture: name exists -> run the recipe and inject the result; name missing -> error.')

h2('6. Same machine, two recipes')
table(['', 'page', 'testimonialsPage'],
      [['Who wrote the recipe?', 'Playwright (built-in)', 'You (in fixtures.ts)'],
       ['What the recipe does', 'builds a browser tab', 'does new TestimonialsPage(page)'],
       ['How you ask for it', '({ page })', '({ testimonialsPage })'],
       ['Who runs the recipe?', 'Playwright', 'Playwright'],
       ['Did you write new?', 'No', 'No']])

h2('7. This IS Dependency Injection (DI)')
para('DI means: you ask for what you need, and the framework provides it — you do not build it '
     'yourself. A dependency = something your code needs (here, a TestimonialsPage). Injection = the '
     'framework gives it to you. Java comparison: this is exactly Spring @Autowired — declare a '
     'dependency, the framework supplies the instance, no new. Playwright fixtures are the same idea '
     'for tests.')

h2('8. The use callback (setup / teardown)')
code("const testimonialsPage = new TestimonialsPage(page);   // 1. SETUP — build it\n"
     "await use(testimonialsPage);                           // 2. hand to test; test RUNS here\n"
     "// 3. TEARDOWN — cleanup would go here")
para('use is not "return the object". It hands the object to the test, pauses the fixture, lets the '
     'test run, then comes back for cleanup. Code before use = setup; code after use = teardown. A '
     'plain return would kill the fixture immediately and lose the teardown phase. (Librarian: '
     'fetches the book, waits while you read, reshelves it.)')

h2('9. Why tests import from fixtures, not @playwright/test')
code("import { test, expect } from '../../utils/fixtures';")
para('base.extend(...) created a NEW test that knows your custom fixtures. The plain test from '
     '@playwright/test would not know testimonialsPage exists. So tests import YOUR test = '
     'Playwright test + your fixtures.')

h2('10. Visual — the recipe machine')
try:
    doc.add_picture('di-flow.png', width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception:
    para('[diagram di-flow.png not found]', italic=True)

h2('The one line to remember')
para('A fixture is a recipe you register so Playwright can build something by name and inject it '
     'into your test — no new. That is Dependency Injection (like Spring @Autowired): you ask, the '
     'framework provides. { testimonialsPage } works because its recipe exists; { bananaPage } fails '
     'because it does not.', italic=True)

# ============================ 8 SPEC ============================
h1('8. Putting It Together: The Spec File')
para('tests/e2e/testimonials.spec.ts — every concept in one place:')
code("// MODULES: import from OUR fixtures (not '@playwright/test')\n"
     "import { test, expect } from '../../utils/fixtures';\n\n"
     "// FUNCTION (arrow) passed to a HOOK; async + await on a browser action\n"
     "test.beforeEach(async ({ page }) => {\n"
     "  await page.goto('/home');                 // logged in already via auth.json\n"
     "});\n\n"
     "// DEPENDENCY INJECTION: declare the fixtures we need\n"
     "test('TC_TST_01 - navigates to Testimonials screen', async ({ testimonialsPage, page }) => {\n"
     "  await testimonialsPage.navigateToTestimonials();         // POM method; awaited\n"
     "  await expect(page).toHaveURL(/framework\\/testimonial/); // auto-retrying matcher -> awaited\n"
     "});")
para('Who creates the TestimonialsPage object, and when? The fixture does, via new, right before the '
     'test runs — because the test declared testimonialsPage as a parameter. The test never calls '
     'new. That is DI in action.')
para('Why import from ../../utils/fixtures and not @playwright/test? Because our fixtures file exports '
     'the extended test that knows our custom fixtures, plus a re-exported expect. Importing from '
     'Playwright directly gives the plain test with no testimonialsPage fixture.')

# ============================ 8B TYPESCRIPT TYPES (PLAIN) ============================
doc.add_page_break()
h1('8B. The TypeScript Type System — Plain & Simple (Revision)')
para('The actual "TS layer" on top of JavaScript. Grounded in config.ts and fixtures.ts.', italic=True)

h2('1. What TypeScript adds (the why)')
para('Plain JavaScript has no type checking. This is legal JS: let age = 25; age = "hello";. JS does '
     'not care, and the mistake blows up later at runtime. TypeScript adds a checker: you label the '
     'type, and TS catches violations while you type, before the code runs.')
code('let age: number = 25;\n'
     'age = "hello";   // TS ERROR: Type \'string\' is not assignable to type \'number\'')
para('This is exactly Java\'s mindset — strongly typed, caught before running. TS brings that safety '
     'on top of JavaScript. JS runs it; TS checks it first.')

h2('2. Type annotations — the basics')
code('let name: string = "Raghav";\nlet age: number = 25;\nlet isActive: boolean = true;')
para('Common types: string, number, boolean, plus Page and Locator (Playwright types). Same as Java '
     'String, int, boolean — just written AFTER the name.')

h2('3. Union types — pick from a fixed set (new vs Java)')
para('From config.ts:')
code("export const ENV = (process.env.ENV ?? 'preprod') as 'dev' | 'preprod' | 'prod' | 'digipulse';")
para('The | means OR. This type says ENV must be ONE of exactly these four strings — nothing else. '
     'Not any string, these four.')
code("let env: 'dev' | 'preprod' | 'prod';\n"
     "env = 'dev';      // OK\n"
     "env = 'staging';  // TS ERROR — not in the list")
para('Why safer than plain string: string allows any text, so a typo like \'prdo\' slips through. The '
     'union allows only the valid values, so TS catches the typo instantly. Java comparison: closest '
     'is an enum — a union of strings is a lighter way to say "only these allowed values".')

h2('4. The type keyword — name a shape')
para('From fixtures.ts:')
code("type MyFixtures = {\n"
     "  pushNotificationPage: PushNotificationPage;\n"
     "  documentLibraryPage:  DocumentLibraryPage;\n"
     "  testimonialsPage:     TestimonialsPage;\n"
     "};")
para('type MyFixtures = { ... } gives a name to a shape: "a MyFixtures is an object with these '
     'properties, each of these types". Reuse the name instead of rewriting the shape. Java '
     'comparison: like a small class/interface describing what fields an object has — but no methods, no new.')

h2('5. Interfaces — the close cousin')
code('interface MyFixtures {\n  testimonialsPage: TestimonialsPage;\n}')
para('interface does almost the same job as type for describing object shapes; they are largely '
     'interchangeable here. This project uses type. Java flag: Java interfaces are contracts for '
     'classes (methods to implement); TS interfaces mostly describe the shape of objects (what '
     'properties exist). Same word, slightly different job.')

h2('6. Generics — the <> brackets')
para('A generic is a type with a blank to fill in, written in <>. It is a container that needs to '
     'know what it holds:')
bullet('Promise<void> = a Promise that holds nothing.')
bullet('Promise<string> = a Promise that holds a string.')
bullet('base.extend<MyFixtures>(...) = extend, and the custom fixtures have the shape MyFixtures.')
para('Analogy: a labeled box. "Box" alone is vague; Box<Shoes> = a box of shoes. The <> fills in '
     'what kind. Java comparison: identical to Java generics — List<String>, Optional<User>, '
     'CompletableFuture<String>. Same syntax, same meaning.')

h2('7. Three operators in config.ts')
code("export const ENV = (process.env.ENV ?? 'preprod') as 'dev' | 'preprod' | 'prod' | 'digipulse';\n"
     "export const USER_EMAIL = process.env.USER_EMAIL!;")
bullet('?? (nullish coalescing) = use the left side, but if it is missing (null/undefined) use the right. A fallback/default.')
bullet('as (type assertion) = "trust me, treat this value as this type". Tells TS the string is actually one of the four.')
bullet('! (non-null assertion) = "trust me, this is not null/undefined". Stops TS warning that it might be missing.')
para('as and ! override the checker — you taking responsibility. Use sparingly.')

h2('The one line to remember')
para('JavaScript has no type checking; TypeScript adds a checker that catches type mistakes before '
     'the code runs — the same safety Java has. Annotate with : type. Union types (\'a\' | \'b\') '
     'limit a value to a fixed set. The type keyword names an object shape. Generics (<>) fill in '
     'what a container holds — identical to Java generics. ?? is a fallback, as and ! override the checker.', italic=True)

# ============================ 9 CHEAT SHEET ============================
h1('9. Java <-> TypeScript Cheat Sheet')
table(['Concept', 'Java', 'TypeScript'],
      [['Variable + type', 'int days', 'days: number (type after name)'],
       ['Method return type', 'String f()', 'f(): string'],
       ['Function as a value', 'not possible directly', 'const f = () => {}'],
       ['Visibility (file)', 'package-private', 'no export = file-private'],
       ['Visibility (class)', 'private', 'private (same)'],
       ['Import a value', 'import class, then Config.ENV', "import { ENV } from './config'"],
       ['Async result type', 'CompletableFuture<T>', 'Promise<T>'],
       ['Wait for async result', 'future.get() (blocks thread)', 'await promise (pauses function only)'],
       ['Constructor name', 'same as class', 'always constructor'],
       ['Dependency injection', 'Spring @Autowired', 'Playwright fixtures'],
       ['this', 'current instance', 'current instance (same)'],
       ['new', 'creates instance + runs ctor', 'identical']])

# ============================ 10 QUIZ ============================
h1('10. Active Recall — Quiz Yourself')
para('Cover the answers. Say each out loud before checking.', italic=True)
qa = [
 ('What does "functions are first-class citizens" mean?',
  'A function is a value — it can be stored in a variable, passed as an argument, or returned from '
  'another function, just like a number or string. Impossible in Java directly.'),
 ('export const ENV vs an un-exported BASE_URLS — what is the difference?',
  'ENV is visible to other files; BASE_URLS has no export, so it is file-private — like a private '
  'field. Encapsulation at the file level.'),
 ('Named import vs default import?',
  'Named: curly braces, exact name. Default: no braces, you choose the name. A file has only one default export.'),
 ('Why does JavaScript need Promises?',
  'JS is single-threaded; it cannot block while waiting. Slow operations return a Promise so the one thread stays free.'),
 ('What is wrong with: const h = page.getHeading(); expect(h).toBe("X");',
  'No await — h holds the Promise object, not the string. toBe compares a Promise to a string and fails. Wrong type.'),
 ('Why does await expect(page).toHaveURL(...) need await but expect(h).toBe("X") does not?',
  'toHaveURL polls the live page and retries (async -> Promise -> await). toBe compares a value already in hand (sync).'),
 ('What does the constructor really do?',
  'Turns a freshly new-ed empty object into a usable one: stores the tab on this.page and registers '
  'locators, once. Without it this.page is undefined and actions crash.'),
 ('Why is it safe to define locators in the constructor before navigating?',
  'Locators are lazy — defining one does not search the DOM; the search happens at action time.'),
 ('What is a fixture, and what Java concept does it match?',
  'Playwright\'s dependency injection — register a page object once, tests receive it by naming it as '
  'a parameter. Same as Spring @Autowired.'),
 ('What does await use(obj) do that return obj cannot?',
  'It hands the object to the test and pauses the fixture around it, enabling teardown after the '
  'test. return ends the fixture immediately, losing teardown.'),
]
for i, (q, a) in enumerate(qa, 1):
    p = doc.add_paragraph()
    p.add_run('Q%d. %s' % (i, q)).bold = True
    ap = doc.add_paragraph()
    ar = ap.add_run('A: ' + a)
    ar.italic = True
    ar.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
foot = doc.add_paragraph()
fr = foot.add_run('Built from the Testimonials-feature session. Companion to PROGRESS.md '
                  '(what was built) and Fixes.md (bugs fixed).')
fr.italic = True
fr.font.size = Pt(9)

doc.save('Concepts-Study-Guide.docx')
print('OK: Concepts-Study-Guide.docx written')
